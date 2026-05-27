from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


SOURCE_DOCX = Path("/Users/shivachoudhry/Downloads/HRM204_Starbucks_Assignment_CA2_Updated (1).docx")
OUTPUT_PDF = Path("/Users/shivachoudhry/Downloads/Ai-VendorHub/document_output/HRM204_Starbucks_Assignment_CA2_Recreated_Final.pdf")


styles = getSampleStyleSheet()
normal = ParagraphStyle(
    "Body",
    parent=styles["BodyText"],
    fontName="Times-Roman",
    fontSize=10.5,
    leading=14,
    alignment=TA_LEFT,
    spaceAfter=6,
)
heading = ParagraphStyle(
    "Heading",
    parent=normal,
    fontName="Times-Bold",
    fontSize=14,
    leading=17,
    spaceBefore=10,
    spaceAfter=8,
)
title = ParagraphStyle(
    "Title",
    parent=heading,
    fontSize=15,
    leading=18,
    alignment=TA_CENTER,
)
cell_style = ParagraphStyle(
    "Cell",
    parent=normal,
    fontSize=9.5,
    leading=12,
    spaceAfter=0,
)
cell_center = ParagraphStyle(
    "CellCenter",
    parent=cell_style,
    alignment=TA_CENTER,
)
cell_bold = ParagraphStyle(
    "CellBold",
    parent=cell_style,
    fontName="Times-Bold",
)
cell_bold_center = ParagraphStyle(
    "CellBoldCenter",
    parent=cell_style,
    fontName="Times-Bold",
    alignment=TA_CENTER,
)


def escape(text):
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def paragraph_text(paragraph):
    return paragraph.text.strip()


def paragraph_image_rel_ids(paragraph):
    ids = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rid = blip.get(qn("r:embed"))
        if rid:
            ids.append(rid)
    return ids


def image_bytes_for_rel(doc_part, rid, zip_file):
    rel = doc_part.rels[rid]
    target = rel.target_ref
    if target.startswith("../"):
        target = target[3:]
    if not target.startswith("word/"):
        target = f"word/{target}"
    return zip_file.read(target)


def make_image(image_bytes):
    image = Image(BytesIO(image_bytes))
    max_width = 6.7 * inch
    max_height = 4.8 * inch
    scale = min(max_width / image.imageWidth, max_height / image.imageHeight, 1)
    image.drawWidth = image.imageWidth * scale
    image.drawHeight = image.imageHeight * scale
    image.hAlign = "CENTER"
    return image


def make_paragraph(text):
    if not text:
        return Spacer(1, 5)

    lowered = text.lower()
    if text.isupper() and len(text) < 80:
        return Paragraph(f"<b>{escape(text)}</b>", title)
    if (
        text.endswith(":")
        or lowered.startswith("question")
        or lowered.startswith("peer rating")
        or lowered.startswith("declaration")
        or lowered.startswith("evaluation")
    ):
        return Paragraph(f"<b>{escape(text)}</b>", heading)
    return Paragraph(escape(text), normal)


def table_data_from_docx_table(table, table_index):
    if table_index == 0:
        return [
            ["Name of the faculty member:", "Dr. Ulfat Andrabi", "Course Code:", "HRM204"],
            ["Academic Task No:", "2", "Course Title:", "Compensation Management"],
            ["Date of Allotment:", "25/04/2026", "Date of Submission:", "10/05/2026"],
            ["Student Roll No:", "12310596", "Student Reg. No:", "_______________"],
            ["Term:", "325262", "Section:", "_______________"],
            ["Max. Marks:", "30", "Marks Obtained:", ""],
        ]
    if table_index == 2:
        return [
            ["Sr. No", "Registration No.", "Name of the Student", "Roll No.", "Peer Rating"],
            ["1", "12310596", "Shiva Choudhry", "56", "10/10"],
            ["2", "12309087", "Ashish Raj", "35", "9.3/10"],
            ["3", "12308406", "Hansraj", "26", "8.2/10"],
            ["4", "12308110", "Abhishek kumar", "22", "8.3/10"],
            ["5", "", "", "", ""],
        ]

    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def make_table(table, table_index):
    raw = table_data_from_docx_table(table, table_index)
    if not raw:
        return Spacer(1, 6)

    col_count = max(len(row) for row in raw)
    normalized = [row + [""] * (col_count - len(row)) for row in raw]

    if table_index == 0:
        widths = [1.65 * inch, 1.75 * inch, 1.65 * inch, 1.75 * inch]
    elif table_index == 2:
        widths = [0.75 * inch, 1.55 * inch, 2.05 * inch, 0.9 * inch, 1.15 * inch]
    else:
        usable = 6.7 * inch
        widths = [usable / col_count] * col_count

    data = []
    for row_index, row in enumerate(normalized):
        pdf_row = []
        for col_index, value in enumerate(row):
            style = cell_style
            if row_index == 0:
                style = cell_bold_center if table_index != 0 else cell_bold
            elif table_index == 0 and col_index in (0, 2):
                style = cell_bold
            elif table_index == 2:
                style = cell_center
            pdf_row.append(Paragraph(escape(value), style))
        data.append(pdf_row)

    table_obj = Table(data, colWidths=widths, repeatRows=1 if table_index not in (0, 1) else 0)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.65, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]
    if table_index != 0:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke))
    if table_index == 2:
        commands.append(("ALIGN", (0, 0), (-1, -1), "CENTER"))

    table_obj.setStyle(TableStyle(commands))
    return KeepTogether([table_obj, Spacer(1, 10)])


def iter_block_items(doc):
    body = doc.element.body
    paragraph_index = 0
    table_index = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", doc.paragraphs[paragraph_index], paragraph_index
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            yield "table", doc.tables[table_index], table_index
            table_index += 1


def build_pdf():
    doc = Document(SOURCE_DOCX)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)

    story = []
    with ZipFile(SOURCE_DOCX) as zip_file:
        for kind, item, index in iter_block_items(doc):
            if kind == "paragraph":
                text = paragraph_text(item)
                rel_ids = paragraph_image_rel_ids(item)
                if text:
                    story.append(make_paragraph(text))
                for rid in rel_ids:
                    try:
                        story.append(make_image(image_bytes_for_rel(doc.part, rid, zip_file)))
                        story.append(Spacer(1, 8))
                    except Exception:
                        continue
            else:
                if index == 2:
                    story.append(Paragraph("<b>Peer Rating Table</b>", heading))
                story.append(make_table(item, index))

    pdf = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="HRM204 Starbucks Assignment CA2 Final",
        author="Shiva Choudhry",
    )
    pdf.build(story)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    build_pdf()
