from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt


SOURCE = Path("/Users/shivachoudhry/Downloads/HRM204_Starbucks_Assignment_CA2_Updated (1).docx")
OUTPUT = Path("/Users/shivachoudhry/Downloads/Ai-VendorHub/document_output/HRM204_Starbucks_Assignment_CA2_Final.docx")


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()
    for table in cell.tables:
        table._element.getparent().remove(table._element)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clone_run_style(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size


def replace_cell_text_preserve_first_run(cell, text):
    first_run = None
    for paragraph in cell.paragraphs:
        if paragraph.runs:
            first_run = paragraph.runs[0]
            break

    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    if first_run is not None:
        clone_run_style(first_run, run)
    else:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document(SOURCE)

# First table: assignment metadata table shown in the screenshot.
metadata = doc.tables[0]
updates = {
    (1, 1): "2",
    (2, 1): "25/04/2026",
    (2, 3): "10/05/2026",
    (3, 3): "_______________",
    (4, 3): "_______________",
}

for (row, col), value in updates.items():
    replace_cell_text_preserve_first_run(metadata.cell(row, col), value)

# Peer rating table: ensure the values match the screenshot exactly.
peer = doc.tables[2]
peer_rows = [
    ["Sr. No", "Registration No.", "Name of the Student", "Roll No.", "Peer Rating"],
    ["1", "12310596", "Shiva Choudhry", "56", "10/10"],
    ["2", "12309087", "Ashish Raj", "35", "9.3/10"],
    ["3", "12308406", "Hansraj", "26", "8.2/10"],
    ["4", "12308110", "Abhishek kumar", "22", "8.3/10"],
    ["5", "", "", "", ""],
]

for row_index, row_values in enumerate(peer_rows):
    for col_index, value in enumerate(row_values):
        set_cell_text(
            peer.cell(row_index, col_index),
            value,
            bold=(row_index == 0),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
